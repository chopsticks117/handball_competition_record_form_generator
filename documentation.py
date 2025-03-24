import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog

SAVE = True

def set_font(doc, font_name='PMingLiU', font_size="11"):
    """Set the font for all text in the document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            #run.font.size = Pt(font_size)
    
    # Also set font for table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.name = font_name
                    #run.font.size = Pt(font_size)

def replace_placeholder_advance(doc, placeholder, replacement):
    """
    Replace a placeholder in the document with a replacement value and set the font to PMingLiU.
    Preserves the rest of the paragraph.
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Split the paragraph into runs
            runs = paragraph.runs
            new_runs = []

            for run in runs:
                if placeholder in run.text:
                    # Split the run into parts before and after the placeholder
                    before, placeholder_text, after = run.text.partition(placeholder)

                    # Add the text before the placeholder
                    if before:
                        new_run = paragraph.add_run(before)
                        new_run.font.name = 'PMingLiU'  # Set font to PMingLiU

                    # Add the replacement text with PMingLiU font
                    new_run = paragraph.add_run(replacement)
                    new_run.font.name = 'PMingLiU'  # Set font to PMingLiU

                    # Add the text after the placeholder
                    if after:
                        new_run = paragraph.add_run(after)
                        new_run.font.name = 'PMingLiU'  # Set font to PMingLiU
                else:
                    # Add the run as is
                    new_run = paragraph.add_run(run.text)
                    new_run.font.name = 'PMingLiU'  # Set font to PMingLiU

            # Clear the original runs
            paragraph._p.clear_content()

            # Add the new runs to the paragraph
            for run in new_runs:
                paragraph._p.append(run._element)


def replace_placeholder(doc, placeholder, replacement):
    """
    Replace a placeholder in the document with a replacement value.
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)                
            
def set_cell_alignment(cell, align=WD_ALIGN_PARAGRAPH.CENTER):
    """
    Set the alignment of text in a table cell.
    align: WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, or WD_ALIGN_PARAGRAPH.JUSTIFY
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = align

def set_cell_vertical_alignment(cell, align="center"):
    """
    Set the vertical alignment of text in a table cell.
    align: "top", "center", or "bottom"
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcValign = OxmlElement('w:vAlign')
    tcValign.set(qn('w:val'), align)
    tcPr.append(tcValign)

# output_folder="/Users/tszhowong/Desktop/Handball/Referee/Application/data"
# output_folder="/Users/johnson/Desktop/Johnson/handball_competition_record_form_generator-main/sdata"
def create_docx_from_template(template_path, match_type, matches, referees):
    # Load the template .docx file
    doc = Document(template_path)
    
    # Set match date to '%d-%m-%Y' form
    match_date_str = matches[0]['日期']
    match_date_obj = datetime.strptime(match_date_str, '%d/%m/%Y')
    match_date = match_date_obj.strftime('%d-%m-%Y')

    # Replace placeholders with actual values
    replace_placeholder(doc, "{{比賽日期}}", match_date)
    replace_placeholder(doc, "{{比賽時間}}", matches[0]['時間'][:4]+"-"+matches[-1]['時間'][-4:])
    replace_placeholder(doc, "{{比賽場地}}", matches[0]['地點'])
    replace_placeholder(doc, "{{賽事類型}}", match_type)
    if match_type == "聯賽" or match_type == "手總盃":
        replace_placeholder(doc, "{{主辦機構}}", "中國香港手球總會")
    else:
        replace_placeholder(doc, "{{主辦機構}}", "中國香港學界體育聯會")

    # Find the table in the document
    match_table = doc.tables[0]  # Assuming the first table is the one we want to update

    # Insert match data into the table
    for i, match in enumerate(matches):
        if i >= len(match_table.rows) - 1:  # Skip the header row and avoid index errors
            break
        row = match_table.rows[i + 1]  # Skip the header row

        # Update cell content
        row.cells[0].text = match['場次']  # 場次
        row.cells[1].text = match['時間']  # 時間
        row.cells[2].text = match['組別']  # 組別
        row.cells[3].text = match['主隊'] # 對賽隊伍: 主隊
        row.cells[5].text = match['客隊'] # 對賽隊伍: 客隊

        # Set alignment for each cell
        for cell in row.cells:
            set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)  # Horizontal alignment
            set_cell_vertical_alignment(cell, "center")  # Vertical alignment

    ref_table = doc.tables[1]  # Assuming the second table is the one we want to update
    
    # Insert referee data into the table
    for i, ref in enumerate(referees):
        row = ref_table.rows[i + 1]  # Skip the header row
        row.cells[2].text = ref['裁判編號']  # 裁判編號
        row.cells[1].text = ref['裁判姓名']  # 裁判姓名

        # Set alignment for each cell
        for cell in row.cells:
            set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)  # Horizontal alignment
            set_cell_vertical_alignment(cell, "center")  # Vertical alignment<end_of_insertio
    
    # Set the font for all text in the document
    set_font(doc, font_name='PMingLiU')

    # Create the output folder if it doesn't exist
    # if not os.path.exists(output_folder):
    #     os.makedirs(output_folder)

    # Generate the output filename based on the current date
    output_folder = filedialog.askdirectory(title="Select Save Location")
    output_filename = f"{match_date}聯賽比賽記錄表.docx"
    output_path = os.path.join(output_folder, output_filename)
    

    # Save the new document
    try:
        if SAVE: # Debug mode
            doc.save(output_path)
        print(f"New document created: {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

def update_txt_file(content, file_path):
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
        print("PDF content saved to output.txt")
    except Exception as e:
        print(f"Error saving PDF content to file: {e}")
